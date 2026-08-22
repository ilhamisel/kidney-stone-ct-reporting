# -*- coding: utf-8 -*-
"""Stage 27 — build a Word document from a Markdown source.

Why a custom builder rather than Pandoc: revised text has to appear in RED, and
Pandoc's docx writer does not carry colour at the run level. Setting the colour
per run through python-docx is the only reliable route.

Supported syntax:
  # / ## / ###        headings
  [[...]]             a RED run, marking text changed in the revision
  **...**             bold
  *...*               italik
  | a | b |           Markdown tablosu -> Word tablosu
  [FIG:file|caption]  a figure with its caption, resolved against the figure dir
  ---                 a horizontal separator, rendered as an empty paragraph

Usage:
  python scripts/27_build_manuscript.py 11_paper/manuscript_revised_EN.md
  python scripts/27_build_manuscript.py 11_paper/response_to_reviewers_EN.md --no-figures
"""
from __future__ import annotations

import argparse
import io
import re
import sys
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))
if __name__ == "__main__":
    try:
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
    except (AttributeError, ValueError):
        pass

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from common.paths import ROOT

FIGDIR = ROOT / "10_examples" / "figures"
RED = RGBColor(0xC0, 0x00, 0x00)
BODY_PT = 11
MAX_FIG_IN = 6.3

TOKEN = re.compile(r"(\[\[.*?\]\]|\*\*.*?\*\*|\*[^*]+\*)", re.S)


def add_runs(par, text: str, red: bool = False):
    """Split text into runs: [[...]] red, **..** bold, *..* italic."""
    for part in TOKEN.split(text):
        if not part:
            continue
        r_red, bold, italic = red, False, False
        if part.startswith("[[") and part.endswith("]]"):
            part, r_red = part[2:-2], True
        elif part.startswith("**") and part.endswith("**"):
            part, bold = part[2:-2], True
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            part, italic = part[1:-1], True
        # a second pass, for nested markers such as [[**x**]]
        inner = TOKEN.split(part)
        if len(inner) > 1:
            for ip in inner:
                if not ip:
                    continue
                b2, i2 = bold, italic
                if ip.startswith("**") and ip.endswith("**"):
                    ip, b2 = ip[2:-2], True
                elif ip.startswith("*") and ip.endswith("*") and len(ip) > 2:
                    ip, i2 = ip[1:-1], True
                run = par.add_run(ip)
                run.bold, run.italic = b2, i2
                if r_red:
                    run.font.color.rgb = RED
            continue
        run = par.add_run(part)
        run.bold, run.italic = bold, italic
        if r_red:
            run.font.color.rgb = RED


def add_table(doc, rows):
    header, body = rows[0], rows[2:]          # rows[1] is the alignment row
    t = doc.add_table(rows=1, cols=len(header))
    t.style = "Table Grid"
    for i, cell in enumerate(header):
        p = t.rows[0].cells[i].paragraphs[0]
        add_runs(p, cell)
        for r in p.runs:
            r.bold = True
            r.font.size = Pt(9)
    for row in body:
        cells = t.add_row().cells
        for i, cell in enumerate(row[:len(header)]):
            p = cells[i].paragraphs[0]
            add_runs(p, cell)
            for r in p.runs:
                r.font.size = Pt(9)
    doc.add_paragraph()


def add_figure(doc, name: str, caption: str) -> bool:
    path = FIGDIR / name
    if not path.exists():
        print(f"  WARNING: figure not found, skipped -> {name}")
        return False
    from PIL import Image
    w, h = Image.open(path).size
    width = min(MAX_FIG_IN, MAX_FIG_IN)
    # tall figures are narrowed by aspect ratio so they fit the page
    if h / w > 1.35:
        width = min(width, 4.6)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_runs(cap, caption, red=True)
    for r in cap.runs:
        r.font.size = Pt(9)
        r.italic = True
    return True


def build(src: Path, out: Path, with_figures: bool) -> tuple[int, int]:
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Times New Roman"
    st.font.size = Pt(BODY_PT)
    st.paragraph_format.space_after = Pt(6)

    lines = src.read_text(encoding="utf-8").split("\n")
    i, n_fig, n_tab = 0, 0, 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        m = re.match(r"^(#{1,3})\s+(.*)$", line)
        if m:
            level = len(m.group(1))
            h = doc.add_heading(level=min(level, 3))
            add_runs(h, m.group(2))
            for r in h.runs:
                r.font.color.rgb = RGBColor(0, 0, 0)
            i += 1
            continue

        m = re.match(r"^\[FIG:([^|]+)\|(.*)\]$", line)
        if m:
            if with_figures and add_figure(doc, m.group(1).strip(), m.group(2).strip()):
                n_fig += 1
            i += 1
            continue

        if line.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append([c.strip() for c in lines[i].strip().strip("|").split("|")])
                i += 1
            if len(rows) >= 3:
                add_table(doc, rows)
                n_tab += 1
            continue

        if line.strip() == "---":
            doc.add_paragraph()
            i += 1
            continue

        if re.match(r"^[-*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, re.sub(r"^[-*]\s+", "", line))
            i += 1
            continue

        if re.match(r"^\d+\.\s+", line):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", line))
            i += 1
            continue

        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        add_runs(p, line)
        i += 1

    doc.save(out)
    return n_fig, n_tab


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("src")
    ap.add_argument("--out", default=None)
    ap.add_argument("--no-figures", action="store_true")
    args = ap.parse_args()
    src = Path(args.src)
    if not src.is_absolute():
        src = ROOT / src
    out = Path(args.out) if args.out else src.with_suffix(".docx")
    n_fig, n_tab = build(src, out, not args.no_figures)
    print(f"wrote: {out}  ({n_fig} figures, {n_tab} tables, "
          f"{out.stat().st_size // 1024} KB)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
